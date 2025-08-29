# api_reports.py

from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from flask_restful import Api, Resource
from typing import Dict, List, Any, Optional
import sqlite3
import json
from pathlib import Path
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io
import base64
from jinja2 import Environment, FileSystemLoader
import threading
import queue
from concurrent.futures import ThreadPoolExecutor


class PotteryAPI:
    """REST API for pottery data access"""
    
    def __init__(self, db_path: str = "pypotterylens.db"):
        self.app = Flask(__name__)
        CORS(self.app)
        self.api = Api(self.app)
        self.db_path = db_path
        
        # Background processing queue
        self.task_queue = queue.Queue()
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        # Setup routes
        self._setup_routes()
    
    def _setup_routes(self):
        """Setup API routes"""
        
        class PotteryItemResource(Resource):
            def __init__(self, db_path):
                self.db_path = db_path
            
            def get(self, item_id=None):
                """Get pottery item(s)"""
                conn = sqlite3.connect(self.db_path)
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                if item_id:
                    cursor.execute("SELECT * FROM pottery_items WHERE id = ?", (item_id,))
                    item = cursor.fetchone()
                    if item:
                        result = dict(item)
                        # Get metadata
                        cursor.execute("SELECT * FROM metadata WHERE pottery_id = ?", (item_id,))
                        metadata = cursor.fetchall()
                        result['metadata'] = {row['key']: row['value'] for row in metadata}
                        conn.close()
                        return jsonify(result)
                    else:
                        conn.close()
                        return {'error': 'Item not found'}, 404
                else:
                    # Get all items with pagination
                    page = request.args.get('page', 1, type=int)
                    per_page = request.args.get('per_page', 20, type=int)
                    offset = (page - 1) * per_page
                    
                    cursor.execute("""
                        SELECT * FROM pottery_items 
                        ORDER BY id DESC 
                        LIMIT ? OFFSET ?
                    """, (per_page, offset))
                    
                    items = [dict(row) for row in cursor.fetchall()]
                    
                    # Get total count
                    cursor.execute("SELECT COUNT(*) as total FROM pottery_items")
                    total = cursor.fetchone()['total']
                    
                    conn.close()
                    
                    return jsonify({
                        'items': items,
                        'total': total,
                        'page': page,
                        'per_page': per_page,
                        'total_pages': (total + per_page - 1) // per_page
                    })
            
            def post(self):
                """Create new pottery item"""
                data = request.get_json()
                
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                
                cursor.execute("""
                    INSERT INTO pottery_items 
                    (filename, source_pdf, source_folder, page_number, instance_number, type, position, rotation)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    data.get('filename'),
                    data.get('source_pdf'),
                    data.get('source_folder'),
                    data.get('page_number'),
                    data.get('instance_number'),
                    data.get('type'),
                    data.get('position'),
                    data.get('rotation')
                ))
                
                item_id = cursor.lastrowid
                conn.commit()
                conn.close()
                
                return {'id': item_id, 'message': 'Item created successfully'}, 201
            
            def put(self, item_id):
                """Update pottery item"""
                data = request.get_json()
                
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                
                # Build update query dynamically
                fields = []
                values = []
                for field in ['type', 'position', 'rotation']:
                    if field in data:
                        fields.append(f"{field} = ?")
                        values.append(data[field])
                
                if fields:
                    values.append(item_id)
                    query = f"UPDATE pottery_items SET {', '.join(fields)}, last_modified = CURRENT_TIMESTAMP WHERE id = ?"
                    cursor.execute(query, values)
                    conn.commit()
                
                conn.close()
                
                return {'message': 'Item updated successfully'}
            
            def delete(self, item_id):
                """Delete pottery item"""
                conn = sqlite3.connect(self.db_path)
                cursor = conn.cursor()
                
                cursor.execute("DELETE FROM pottery_items WHERE id = ?", (item_id,))
                cursor.execute("DELETE FROM metadata WHERE pottery_id = ?", (item_id,))
                cursor.execute("DELETE FROM annotations WHERE pottery_id = ?", (item_id,))
                
                conn.commit()
                conn.close()
                
                return {'message': 'Item deleted successfully'}
        
        class SearchResource(Resource):
            def __init__(self, db_path):
                self.db_path = db_path
            
            def get(self):
                """Search pottery items"""
                query = request.args.get('q', '')
                filters = request.args.to_dict()
                filters.pop('q', None)
                
                conn = sqlite3.connect(self.db_path)
                conn.row_factory = sqlite3.Row
                cursor = conn.cursor()
                
                # Build search query
                conditions = []
                params = []
                
                if query:
                    conditions.append("(filename LIKE ? OR type LIKE ? OR position LIKE ?)")
                    params.extend([f'%{query}%'] * 3)
                
                for field, value in filters.items():
                    if field in ['type', 'position', 'rotation']:
                        conditions.append(f"{field} = ?")
                        params.append(value)
                
                where_clause = " AND ".join(conditions) if conditions else "1=1"
                
                cursor.execute(f"""
                    SELECT * FROM pottery_items 
                    WHERE {where_clause}
                    ORDER BY id DESC
                """, params)
                
                results = [dict(row) for row in cursor.fetchall()]
                conn.close()
                
                return jsonify({
                    'results': results,
                    'count': len(results),
                    'query': query,
                    'filters': filters
                })
        
        class BatchProcessResource(Resource):
            def __init__(self, db_path, task_queue, executor):
                self.db_path = db_path
                self.task_queue = task_queue
                self.executor = executor
            
            def post(self):
                """Submit batch processing job"""
                data = request.get_json()
                task_type = data.get('task_type')
                items = data.get('items', [])
                options = data.get('options', {})
                
                # Generate task ID
                task_id = f"task_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(str(data)) % 10000}"
                
                # Submit task to background processing
                future = self.executor.submit(
                    self._process_batch_task,
                    task_id, task_type, items, options
                )
                
                # Store task info
                self.task_queue.put({
                    'id': task_id,
                    'future': future,
                    'status': 'processing',
                    'created': datetime.now()
                })
                
                return {
                    'task_id': task_id,
                    'message': 'Batch processing started',
                    'status_url': f'/api/batch/status/{task_id}'
                }, 202
            
            def _process_batch_task(self, task_id, task_type, items, options):
                """Process batch task in background"""
                # This is where you would implement actual batch processing
                # For now, just a placeholder
                import time
                time.sleep(5)  # Simulate processing
                
                return {
                    'task_id': task_id,
                    'status': 'completed',
                    'processed': len(items),
                    'results': 'Batch processing completed'
                }
        
        # Register resources
        self.api.add_resource(
            PotteryItemResource, 
            '/api/items', 
            '/api/items/<int:item_id>',
            resource_class_kwargs={'db_path': self.db_path}
        )
        
        self.api.add_resource(
            SearchResource,
            '/api/search',
            resource_class_kwargs={'db_path': self.db_path}
        )
        
        self.api.add_resource(
            BatchProcessResource,
            '/api/batch',
            resource_class_kwargs={
                'db_path': self.db_path,
                'task_queue': self.task_queue,
                'executor': self.executor
            }
        )
        
        # Add export endpoints
        @self.app.route('/api/export/<format>')
        def export_data(format):
            """Export data in various formats"""
            if format == 'cidoc-crm':
                # Export CIDOC-CRM
                from cidoc_crm_export import CIDOCCRMExporter
                exporter = CIDOCCRMExporter()
                # Implementation here
                return jsonify({'message': 'Export initiated'})
            
            elif format == 'geojson':
                # Export GeoJSON
                from image_processing_advanced import GISExporter
                exporter = GISExporter()
                # Implementation here
                return jsonify({'message': 'Export initiated'})
            
            else:
                return jsonify({'error': 'Unknown export format'}), 400
    
    def run(self, host='0.0.0.0', port=5000, debug=False):
        """Run the API server"""
        self.app.run(host=host, port=port, debug=debug)


class ReportGenerator:
    """Generates reports in multiple formats and languages"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.languages = {
            'en': {
                'title': 'Pottery Documentation Report',
                'summary': 'Summary',
                'typology': 'Typology',
                'measurements': 'Measurements',
                'context': 'Archaeological Context',
                'bibliography': 'Bibliography',
                'date_generated': 'Date Generated'
            },
            'it': {
                'title': 'Rapporto di Documentazione Ceramica',
                'summary': 'Riassunto',
                'typology': 'Tipologia',
                'measurements': 'Misure',
                'context': 'Contesto Archeologico',
                'bibliography': 'Bibliografia',
                'date_generated': 'Data di Generazione'
            },
            'es': {
                'title': 'Informe de Documentación Cerámica',
                'summary': 'Resumen',
                'typology': 'Tipología',
                'measurements': 'Medidas',
                'context': 'Contexto Arqueológico',
                'bibliography': 'Bibliografía',
                'date_generated': 'Fecha de Generación'
            }
        }
    
    def generate_pdf_report(self, data: Dict, output_path: str, 
                           template: str = 'standard', language: str = 'en'):
        """Generate PDF report with custom template"""
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        trans = self.languages.get(language, self.languages['en'])
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_CENTER
        )
        story.append(Paragraph(trans['title'], title_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Date
        date_style = ParagraphStyle(
            'DateStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#7f8c8d'),
            alignment=TA_CENTER
        )
        story.append(Paragraph(
            f"{trans['date_generated']}: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            date_style
        ))
        story.append(Spacer(1, 0.5*inch))
        
        # Summary section
        if 'summary' in data:
            story.append(Paragraph(trans['summary'], self.styles['Heading1']))
            story.append(Paragraph(data['summary'], self.styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
        
        # Items table
        if 'items' in data:
            table_data = [['ID', trans['typology'], trans['measurements'], 'Status']]
            
            for item in data['items']:
                row = [
                    item.get('id', ''),
                    item.get('type', ''),
                    f"H: {item.get('height', 'N/A')} mm",
                    item.get('condition', '')
                ]
                table_data.append(row)
            
            table = Table(table_data, colWidths=[1*inch, 2*inch, 2*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            story.append(PageBreak())
        
        # Images section
        if 'images' in data:
            for img_data in data['images']:
                if 'path' in img_data:
                    img = RLImage(img_data['path'], width=4*inch, height=3*inch)
                    story.append(img)
                    
                    if 'caption' in img_data:
                        caption_style = ParagraphStyle(
                            'Caption',
                            parent=self.styles['Normal'],
                            fontSize=10,
                            textColor=colors.HexColor('#7f8c8d'),
                            alignment=TA_CENTER
                        )
                        story.append(Paragraph(img_data['caption'], caption_style))
                    story.append(Spacer(1, 0.3*inch))
        
        # Bibliography
        if 'bibliography' in data:
            story.append(Paragraph(trans['bibliography'], self.styles['Heading1']))
            for ref in data['bibliography']:
                story.append(Paragraph(f"• {ref}", self.styles['Normal']))
        
        # Build PDF
        doc.build(story)
    
    def generate_docx_report(self, data: Dict, output_path: str, 
                           template: str = 'standard', language: str = 'en'):
        """Generate DOCX report"""
        doc = Document()
        trans = self.languages.get(language, self.languages['en'])
        
        # Title
        title = doc.add_heading(trans['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(
            f"{trans['date_generated']}: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Summary
        if 'summary' in data:
            doc.add_heading(trans['summary'], level=1)
            doc.add_paragraph(data['summary'])
        
        # Items table
        if 'items' in data:
            doc.add_heading('Items', level=1)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Shading Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = trans['typology']
            hdr_cells[2].text = trans['measurements']
            hdr_cells[3].text = 'Status'
            
            # Data rows
            for item in data['items']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('id', ''))
                row_cells[1].text = item.get('type', '')
                row_cells[2].text = f"H: {item.get('height', 'N/A')} mm"
                row_cells[3].text = item.get('condition', '')
        
        # Images
        if 'images' in data:
            doc.add_heading('Documentation', level=1)
            for img_data in data['images']:
                if 'path' in img_data and Path(img_data['path']).exists():
                    doc.add_picture(img_data['path'], width=Inches(4))
                    if 'caption' in img_data:
                        caption = doc.add_paragraph(img_data['caption'])
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.runs[0].italic = True
        
        # Save document
        doc.save(output_path)
    
    def generate_web_report(self, data: Dict, output_dir: str, 
                          template: str = 'standard', language: str = 'en'):
        """Generate static website for pottery documentation"""
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # Create directory structure
        (output_path / 'css').mkdir(exist_ok=True)
        (output_path / 'js').mkdir(exist_ok=True)
        (output_path / 'images').mkdir(exist_ok=True)
        
        # Generate CSS
        css_content = """
        body {
            font-family: 'Arial', sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 0;
            background-color: #f4f4f4;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            background-color: white;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        
        h1 {
            color: #2c3e50;
            text-align: center;
            margin-bottom: 30px;
        }
        
        .pottery-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(300px, 1fr));
            gap: 20px;
            margin-top: 30px;
        }
        
        .pottery-item {
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            transition: transform 0.3s;
        }
        
        .pottery-item:hover {
            transform: translateY(-5px);
            box-shadow: 0 5px 15px rgba(0,0,0,0.2);
        }
        
        .pottery-image {
            width: 100%;
            height: 200px;
            object-fit: contain;
            border-radius: 4px;
            margin-bottom: 10px;
        }
        
        .pottery-details {
            font-size: 14px;
            color: #666;
        }
        
        .stats-section {
            margin: 40px 0;
            padding: 20px;
            background-color: #ecf0f1;
            border-radius: 8px;
        }
        """
        
        with open(output_path / 'css' / 'style.css', 'w') as f:
            f.write(css_content)
        
        # Generate HTML
        trans = self.languages.get(language, self.languages['en'])
        
        html_content = f"""
        <!DOCTYPE html>
        <html lang="{language}">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{trans['title']}</title>
            <link rel="stylesheet" href="css/style.css">
        </head>
        <body>
            <div class="container">
                <h1>{trans['title']}</h1>
                <p class="date">{trans['date_generated']}: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
                
                <div class="stats-section">
                    <h2>{trans['summary']}</h2>
                    <p>{data.get('summary', '')}</p>
                    <p>Total items: {len(data.get('items', []))}</p>
                </div>
                
                <div class="pottery-grid">
        """
        
        # Add items
        for item in data.get('items', []):
            html_content += f"""
                    <div class="pottery-item">
                        <img class="pottery-image" src="{item.get('image', 'placeholder.png')}" 
                             alt="{item.get('type', 'Pottery')}">
                        <div class="pottery-details">
                            <h3>{item.get('type', 'Unknown')}</h3>
                            <p><strong>ID:</strong> {item.get('id', 'N/A')}</p>
                            <p><strong>{trans['measurements']}:</strong> 
                               H: {item.get('height', 'N/A')} mm, 
                               D: {item.get('diameter', 'N/A')} mm</p>
                            <p><strong>Context:</strong> {item.get('context', 'N/A')}</p>
                        </div>
                    </div>
            """
        
        html_content += """
                </div>
            </div>
        </body>
        </html>
        """
        
        with open(output_path / 'index.html', 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def generate_bibliography(self, items: List[Dict], format: str = 'chicago') -> List[str]:
        """Generate bibliography entries in various formats"""
        bibliography = []
        
        for item in items:
            if format == 'chicago':
                # Chicago style format
                entry = f"{item.get('author', 'Unknown')}. "
                entry += f"\"{item.get('title', 'Untitled')}.\" "
                entry += f"{item.get('journal', 'Unknown Journal')} "
                entry += f"{item.get('volume', '')}, no. {item.get('issue', '')} "
                entry += f"({item.get('year', 'n.d.')}): {item.get('pages', 'n.p.')}."
                
            elif format == 'apa':
                # APA style format
                entry = f"{item.get('author', 'Unknown')} "
                entry += f"({item.get('year', 'n.d.')}). "
                entry += f"{item.get('title', 'Untitled')}. "
                entry += f"{item.get('journal', 'Unknown Journal')}, "
                entry += f"{item.get('volume', '')}({item.get('issue', '')}), "
                entry += f"{item.get('pages', 'n.p.')}."
            
            bibliography.append(entry)
        
        return bibliography